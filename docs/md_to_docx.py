"""
Markdown to DOCX Converter for Cherag Documentation
Combines multiple markdown files into a single Word document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
import os

def create_styles(doc):
    """Create custom styles for the document"""
    styles = doc.styles
    
    # Modify Heading 1
    h1_style = styles['Heading 1']
    h1_style.font.size = Pt(24)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 51, 102)
    
    # Modify Heading 2
    h2_style = styles['Heading 2']
    h2_style.font.size = Pt(18)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 76, 153)
    
    # Modify Heading 3
    h3_style = styles['Heading 3']
    h3_style.font.size = Pt(14)
    h3_style.font.bold = True
    
    return doc

def parse_markdown_line(doc, line, in_code_block, code_language):
    """Parse a single markdown line and add to document"""
    line = line.rstrip()
    
    # Check for code block markers
    if line.startswith('```'):
        if in_code_block:
            return False, None  # End code block
        else:
            lang = line[3:].strip()
            return True, lang  # Start code block
    
    # If in code block, add as code
    if in_code_block:
        para = doc.add_paragraph()
        run = para.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        para.paragraph_format.left_indent = Inches(0.5)
        return True, code_language
    
    # Empty line
    if not line:
        return False, None
    
    # Headings
    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('#### '):
        doc.add_heading(line[5:], level=4)
    elif line.startswith('##### '):
        doc.add_heading(line[6:], level=5)
    # Horizontal rule
    elif line.startswith('---') or line.startswith('***'):
        doc.add_paragraph('─' * 50)
    # Bullet points
    elif line.startswith('- ') or line.startswith('* '):
        para = doc.add_paragraph(style='List Bullet')
        add_formatted_text(para, line[2:])
    elif line.startswith('  - ') or line.startswith('  * '):
        para = doc.add_paragraph(style='List Bullet 2')
        add_formatted_text(para, line[4:])
    # Numbered lists
    elif re.match(r'^\d+\.\s', line):
        match = re.match(r'^\d+\.\s(.*)$', line)
        para = doc.add_paragraph(style='List Number')
        add_formatted_text(para, match.group(1))
    # Table rows (simplified handling)
    elif line.startswith('|'):
        # Skip table separator rows
        if re.match(r'^\|[\s\-:]+\|', line):
            pass
        else:
            para = doc.add_paragraph()
            # Clean up table formatting
            cells = [c.strip() for c in line.split('|')[1:-1]]
            add_formatted_text(para, '  |  '.join(cells))
    # Blockquotes
    elif line.startswith('> '):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.right_indent = Inches(0.5)
        add_formatted_text(para, line[2:])
    # Regular paragraph
    else:
        para = doc.add_paragraph()
        add_formatted_text(para, line)
    
    return False, None

def add_formatted_text(para, text):
    """Add text with inline formatting (bold, italic, code)"""
    # Simple pattern matching for formatting
    # This handles **bold**, *italic*, and `code`
    
    i = 0
    while i < len(text):
        # Check for bold
        if text[i:i+2] == '**':
            end = text.find('**', i+2)
            if end != -1:
                run = para.add_run(text[i+2:end])
                run.bold = True
                i = end + 2
                continue
        
        # Check for code
        if text[i] == '`':
            end = text.find('`', i+1)
            if end != -1:
                run = para.add_run(text[i+1:end])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                i = end + 1
                continue
        
        # Check for italic
        if text[i] == '*' and (i == 0 or text[i-1] != '*'):
            end = text.find('*', i+1)
            if end != -1 and (end+1 >= len(text) or text[end+1] != '*'):
                run = para.add_run(text[i+1:end])
                run.italic = True
                i = end + 1
                continue
        
        # Find next special character
        next_special = len(text)
        for char in ['**', '`', '*']:
            pos = text.find(char, i+1)
            if pos != -1 and pos < next_special:
                next_special = pos
        
        # Add regular text up to next special character
        run = para.add_run(text[i:next_special])
        i = next_special

def convert_markdown_to_docx(md_files, output_path):
    """Convert multiple markdown files to a single DOCX"""
    doc = Document()
    doc = create_styles(doc)
    
    # Add title page
    title = doc.add_heading('Cherág - AI Study Partner', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Complete Software Engineering Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    date_para = doc.add_paragraph('Generated: January 2026')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Add table of contents placeholder
    toc_heading = doc.add_heading('Table of Contents', 1)
    doc.add_paragraph('1. Software Documentation')
    doc.add_paragraph('2. API Reference')
    doc.add_paragraph('3. Database Design')
    doc.add_paragraph('4. User Guide')
    doc.add_page_break()
    
    # Process each markdown file
    for md_file in md_files:
        print(f"Processing: {md_file}")
        
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        in_code_block = False
        code_language = None
        
        for line in content.split('\n'):
            in_code_block, code_language = parse_markdown_line(
                doc, line, in_code_block, code_language
            )
        
        # Add page break between sections
        doc.add_page_break()
    
    # Save document
    doc.save(output_path)
    print(f"\nDocument saved to: {output_path}")

if __name__ == '__main__':
    docs_dir = r'c:\Users\HALA-MADRID\Desktop\Cherag\docs'
    
    md_files = [
        os.path.join(docs_dir, 'SOFTWARE_DOCUMENTATION.md'),
        os.path.join(docs_dir, 'API_REFERENCE.md'),
        os.path.join(docs_dir, 'DATABASE_DESIGN.md'),
        os.path.join(docs_dir, 'USER_GUIDE.md'),
    ]
    
    output_path = os.path.join(docs_dir, 'Cherag_Complete_Documentation.docx')
    
    convert_markdown_to_docx(md_files, output_path)
